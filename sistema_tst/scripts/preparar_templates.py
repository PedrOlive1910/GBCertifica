"""Prepara os modelos DOCX usados pelo motor de documentos.

O script copia os documentos fornecidos pelo cliente e troca apenas os dados
variáveis por tags Jinja compatíveis com ``docxtpl``. A estrutura, as imagens,
as assinaturas e a formatação original são preservadas.
"""

from __future__ import annotations

import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from lxml import etree


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = PROJECT_ROOT.parent
UPLOAD_ROOT = WORKSPACE_ROOT / "upload"
OUTPUT_ROOT = PROJECT_ROOT / "app" / "document_templates"


def replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    """Substitui texto mesmo quando o Word o dividiu em vários runs."""
    if not old or old not in paragraph.text:
        return False

    while old in paragraph.text:
        full_text = "".join(run.text for run in paragraph.runs)
        start = full_text.find(old)
        if start < 0:
            break
        end = start + len(old)

        positions: list[tuple[int, int]] = []
        offset = 0
        for run_index, run in enumerate(paragraph.runs):
            positions.append((offset, offset + len(run.text)))
            offset += len(run.text)

        first_index = next(
            index for index, (_, run_end) in enumerate(positions) if run_end > start
        )
        last_index = next(
            index
            for index, (run_start, run_end) in enumerate(positions)
            if run_start < end <= run_end
        )

        first_run = paragraph.runs[first_index]
        first_start, _ = positions[first_index]
        _, last_end = positions[last_index]
        prefix = first_run.text[: start - first_start]
        suffix = paragraph.runs[last_index].text[
            len(paragraph.runs[last_index].text) - (last_end - end) :
        ]

        first_run.text = prefix + new + suffix
        for index in range(first_index + 1, last_index + 1):
            paragraph.runs[index].text = ""

    return True


def replace_everywhere(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in iter_paragraphs(document):
        for old, new in replacements.items():
            replace_in_paragraph(paragraph, old, new)


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from iter_table_paragraphs(table)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table: Table):
    visited = set()
    for row in table.rows:
        for cell in row.cells:
            cell_element = cell._tc
            if cell_element in visited:
                continue
            visited.add(cell_element)
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def set_cell_text(cell: _Cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def unique_cell(row, column: int) -> _Cell:
    return row.cells[column]


def remove_row(table: Table, row_index: int) -> None:
    row = table.rows[row_index]
    row._element.getparent().remove(row._element)


def replace_in_xml_paragraph(paragraph, old: str, new: str) -> bool:
    nodes = paragraph.xpath('.//*[local-name()="t"]')
    if not nodes:
        return False
    changed = False
    while True:
        full_text = "".join(node.text or "" for node in nodes)
        start = full_text.find(old)
        if start < 0:
            break
        end = start + len(old)
        offsets = []
        cursor = 0
        for node in nodes:
            value = node.text or ""
            offsets.append((cursor, cursor + len(value)))
            cursor += len(value)
        first = next(index for index, (_, node_end) in enumerate(offsets) if node_end > start)
        last = next(index for index, (node_start, node_end) in enumerate(offsets) if node_start < end <= node_end)
        first_start, _ = offsets[first]
        _, last_end = offsets[last]
        prefix = (nodes[first].text or "")[: start - first_start]
        suffix_value = nodes[last].text or ""
        suffix = suffix_value[len(suffix_value) - (last_end - end) :]
        nodes[first].text = prefix + new + suffix
        for index in range(first + 1, last + 1):
            nodes[index].text = ""
        changed = True
    return changed


def patch_docx_xml(path: Path, replacements: dict[str, str]) -> None:
    """Aplica tags também em caixas de texto, WordArt e formas."""
    temporary = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as source, ZipFile(temporary, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                except etree.XMLSyntaxError:
                    root = None
                if root is not None:
                    changed = False
                    for paragraph in root.xpath('//*[local-name()="p"]'):
                        for old, new in replacements.items():
                            changed |= replace_in_xml_paragraph(paragraph, old, new)
                    if changed:
                        data = etree.tostring(
                            root,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone=True,
                        )
            target.writestr(info, data)
    temporary.replace(path)


def remove_trailing_empty_paragraphs(path: Path) -> None:
    """Evita uma página vazia após a tabela de página inteira da OS."""
    temporary = path.with_suffix(".trimmed.docx")
    with ZipFile(path, "r") as source, ZipFile(temporary, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data)
                body = root.xpath('//*[local-name()="body"]')[0]
                children = list(body)
                for child in reversed(children):
                    tag = etree.QName(child).localname
                    if tag == "sectPr":
                        continue
                    if tag != "p":
                        break
                    has_content = child.xpath(
                        './/*[local-name()="t" or local-name()="drawing" or local-name()="pict" or local-name()="br"]'
                    )
                    if has_content:
                        break
                    body.remove(child)
                for section_type in body.xpath(
                    './*[local-name()="sectPr"]/*[local-name()="type"]'
                ):
                    section_type.getparent().remove(section_type)
                word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                for margins in body.xpath(
                    './*[local-name()="sectPr"]/*[local-name()="pgMar"]'
                ):
                    margins.set(f"{{{word_namespace}}}bottom", "0")
                for row in body.xpath('.//*[local-name()="tr"]'):
                    row_text = "".join(
                        row.xpath('.//*[local-name()="t"]/text()')
                    )
                    if "{{ epis_atividade_texto }}" not in row_text:
                        continue
                    for height in row.xpath(
                        './*[local-name()="trPr"]/*[local-name()="trHeight"]'
                    ):
                        height.set(f"{{{word_namespace}}}val", "2800")
                data = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            target.writestr(info, data)
    temporary.replace(path)


def preparar_nr06(source: Path, destination: Path) -> None:
    shutil.copy2(source, destination)
    patch_docx_xml(destination, {"{{ data_emissao }}": "{{ data_documento_extenso }}"})


def preparar_nr12(source: Path, destination: Path) -> None:
    shutil.copy2(source, destination)
    patch_docx_xml(
        destination,
        {
            "FABIO CARVALHO DA SILVA": "{{ nome_funcionario }}",
            "096.237.243-90": "{{ cpf_funcionario }}",
            "Furadeira, Martelete, Lixadeira, Compactador manual de solo,Serra Circular manual,Esmerilhadeira": "{{ maquinas_texto }}",
            "MORAIS CONSTRULOK CONSTRUÇÕES E LOCAÇÕES LTDA": "{{ razao_social }}",
            "MORAIS CONSTRULOK CONSTRUÇÕES LTDA": "{{ razao_social }}",
            "44.413.036/0001-72": "{{ cnpj }}",
            "Rua Antonio Benicio de Queiroz, 62 – Bairro Redenção – Monte Custódia/ PE, CEP: 56640-000": "{{ endereco_completo }}",
            "Rua Antonio Benicio de Queiroz, 62  – Bairro Redenção – Monte Custódia/ PE, CEP: 56640-000": "{{ endereco_completo }}",
            "Custódia – PE, 13 de Maio de 2026": "{{ cidade_empresa }}, {{ data_documento_extenso }}",
        },
    )


def preparar_nr18(source: Path, destination: Path) -> None:
    shutil.copy2(source, destination)
    patch_docx_xml(
        destination,
        {
            "FABIO CARVALHO DA SILVA": "{{ nome_funcionario }}",
            "096.237.243-90": "{{ cpf_funcionario }}",
            "Morais Construlok Construções e locações LTDA": "{{ razao_social }}",
            "MORAIS CONSTRULOK CONSTRUÇÕES LTDA": "{{ razao_social }}",
            "44.413.036/0001-72": "{{ cnpj }}",
            "Rua Antonio Benicio de Queiroz, 62 – Bairro Redenção – Custódia -PE, CEP: 5640-000": "{{ endereco_completo }}",
            "Rua Antonio Benicio de Queiroz, 62 – Bairro Redenção –  Custódia -PE, CEP: 5640-000": "{{ endereco_completo }}",
            "Custódia PE, 12 de Maio de 2026": "{{ cidade_empresa }}, {{ data_documento_extenso }}",
        },
    )


def preparar_nr35(source: Path, destination: Path) -> None:
    shutil.copy2(source, destination)
    patch_docx_xml(
        destination,
        {
            "FABIO CARVALHO DA SILVA": "{{ nome_funcionario }}",
            "096.237.243-90": "{{ cpf_funcionario }}",
            "MORAIS CONSTRULOK CONSTRUÇÕES E LOCAÇÕES LTDA": "{{ razao_social }}",
            "MORAIS CONSTRULOK CONSTRUÇÕES  LTDA": "{{ razao_social }}",
            "MORAIS CONSTRULOK CONSTRUÇÕES LTDA": "{{ razao_social }}",
            "44.413.036/0001-72": "{{ cnpj }}",
            "Rua Antonio Benicio de Queiroz, 62 – Bairro Redenção – Custódia - PE, CEP: 56640-000": "{{ endereco_completo }}",
            "Rua Antonio Benicio de Queiroz, 62 – Bairro Redenção –  Custódia - PE, CEP: 56640-000": "{{ endereco_completo }}",
            "Rua Antonio Benicio de Queiroz, 62 – Bairro Redenção – CUSTODIA/PE, CEP: 5640-000.": "{{ endereco_completo }}",
            "Rua Antonio Benicio de Queiroz, 62 – RedenÇãO – Custódia – PE, CEP: 56640-000.": "{{ endereco_completo }}",
            "Rua Antonio Benicio de Queiroz, 62 – RedenÇãO –  Custódia – PE, CEP: 56640-000.": "{{ endereco_completo }}",
            "Custodia-PE, 14 de Maio de 2026": "{{ cidade_empresa }}, {{ data_documento_extenso }}",
            "Custódia PE, 14 de Maio de 2026": "{{ cidade_empresa }}, {{ data_documento_extenso }}",
        },
    )


def preparar_ficha_epi(source: Path, destination: Path) -> None:
    document = Document(source)
    replace_everywhere(
        document,
        {
            "FIGUIREDO E SIQUEIRA CONSTRUÇÕES LTDA": "{{ razao_social }}",
            "14.146.041/0001-03": "{{ cnpj }}",
            "ROMERIO REZENDE SÁ": "{{ nome_funcionario }}",
            "Carpinteiro": "{{ funcao }}",
            "10/03/2026": "{{ data_admissao }}",
        },
    )

    epi_table = document.tables[1]
    # A primeira linha permanece como cabeçalho. As três linhas seguintes são
    # o bloco repetível do docxtpl; as demais linhas fixas são removidas.
    set_cell_text(epi_table.rows[1].cells[0], "{%tr for epi in epis %}")
    for cell in epi_table.rows[1].cells[1:]:
        set_cell_text(cell, "")

    values = (
        "{{ epi.quantidade }}",
        "{{ epi.descricao }}",
        "{{ epi.ca }}",
        "{{ epi.data_entrega }}",
        "{{ epi.data_devolucao }}",
        "{{ epi.assinatura }}",
    )
    for cell, value in zip(epi_table.rows[2].cells, values):
        set_cell_text(cell, value)

    set_cell_text(epi_table.rows[3].cells[0], "{%tr endfor %}")
    for cell in epi_table.rows[3].cells[1:]:
        set_cell_text(cell, "")

    while len(epi_table.rows) > 4:
        remove_row(epi_table, 4)

    document.save(destination)


def preparar_ordem_servico(source: Path, destination: Path) -> None:
    document = Document(source)
    table = document.tables[0]

    replace_everywhere(
        document,
        {
            "MORAIS CONSTRULOK": "{{ razao_social }}",
            "444444444444": "{{ cnpj }}",
            "ARMADOR": "{{ funcao }}",
            "OBRA": "{{ setor }}",
            "7153-15": "{{ cbo }}",
            "Exposição a Ruído continuo e intermitente; Calor;": "{{ risco_fisico }}",
            "Inalação de poeiras respiráveis e minerais.": "{{ risco_quimico }}",
            "Não Identificado.": "{{ risco_biologico }}",
            "Posturas incômodas ou pouco confortáveis.": "{{ risco_ergonomico }}",
            "Queda de altura acima de dois metros. Corte e ou perfuração equipamentos rotativos. Queda de mesmo Nivel .": "{{ risco_acidentes }}",
            "FRANCISCO DAS CHAGAS NASCIMENTO MORAES": "{{ nome_funcionario }}",
            "27/08/2024": "{{ data_admissao }}",
            "Elton Bruno A. Olegario": "{{ responsavel_nome }}",
            "INSTRUTOR E TECNICO DE SEGURANÇA DO TRABALHO": "{{ responsavel_cargo }}",
            "0018863/PE": "{{ responsavel_registro }}",
        },
    )
    remove_trailing_empty_paragraphs(destination)

    set_cell_text(unique_cell(table.rows[3], 0), "{{ descricao_funcao }}")
    set_cell_text(unique_cell(table.rows[11], 0), "{{ epis_atividade_texto }}")

    recommendations = unique_cell(table.rows[12], 0)
    recommendation_text = recommendations.text
    prefix = "5. RECOMENDAÇÕES:"
    if prefix in recommendation_text:
        set_cell_text(recommendations, f"{prefix} {{{{ recomendacoes }}}}")

    procedures = unique_cell(table.rows[13], 0)
    procedures_text = procedures.text
    prefix = "6. PROCEDIMENTOS EM CASO DE ACIDENTES:"
    if prefix in procedures_text:
        set_cell_text(procedures, f"{prefix} {{{{ procedimentos_acidente }}}}")

    # A data da assinatura pode ser diferente da admissão.
    replace_in_paragraph(table.rows[16].cells[7].paragraphs[0], "{{ data_admissao }}", "{{ data_documento }}")
    document.save(destination)
    patch_docx_xml(
        destination,
        {
            "Preparam a confecção de armações e estruturas de concreto e de corpos de prova. cortam e dobram ferragens de lajes. montam e aplicam armações de fundações, pilares e vigas. moldam corpos de prova.": "",
            "Exposição a Ruído continuo e intermitente; Calor;": "{{ risco_fisico }}",
            "Inalação de poeiras respiráveis e minerais.": "{{ risco_quimico }}",
            "Não Identificado.": "{{ risco_biologico }}",
            "Posturas incômodas ou pouco confortáveis.": "{{ risco_ergonomico }}",
            "Queda de altura acima de dois metros. Corte e ou perfuração equipamentos rotativos. Queda de mesmo Nivel .": "{{ risco_acidentes }}",
            "FRANCISCO DAS CHAGAS NASCIMENTO MORAES": "{{ nome_funcionario }}",
            "27/08/2024": "{{ data_documento }}",
            "Elton Bruno A. Olegario": "{{ responsavel_nome }}",
            "INSTRUTOR E TECNICO DE SEGURANÇA DO TRABALHO": "{{ responsavel_cargo }}",
            "0018863/PE": "{{ responsavel_registro }}",
        },
    )


def main() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    converted_os = WORKSPACE_ROOT / "doc_review" / "converted" / "Ordem de serviço - ARMADOR (1).docx"

    jobs = (
        (preparar_nr06, UPLOAD_ROOT / "template_nr06(1).docx", OUTPUT_ROOT / "nr06.docx"),
        (preparar_nr12, UPLOAD_ROOT / "Certificado NR 12 Morais CONSTRULOK (1)(1).docx", OUTPUT_ROOT / "nr12.docx"),
        (preparar_nr18, UPLOAD_ROOT / "Certificado NR 18 Morais CONSTRULOK(1).docx", OUTPUT_ROOT / "nr18.docx"),
        (preparar_nr35, UPLOAD_ROOT / "Certificado NR 35 Morais CONSTRULOK(1).docx", OUTPUT_ROOT / "nr35.docx"),
        (preparar_ficha_epi, UPLOAD_ROOT / "Ficha de Epi carpinteiro (1)(1).docx", OUTPUT_ROOT / "ficha_epi.docx"),
        (preparar_ordem_servico, converted_os, OUTPUT_ROOT / "ordem_servico.docx"),
    )

    for builder, source, destination in jobs:
        if not source.exists():
            raise FileNotFoundError(source)
        builder(source, destination)
        print(f"Criado: {destination.relative_to(PROJECT_ROOT)}")

if __name__ == "__main__":
    main()
